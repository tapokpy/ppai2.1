from docx import Document
from openpyxl import load_workbook

from app.services.doc_generator import generate_docx, generate_xlsx


def test_generate_docx_creates_file_with_expected_content():
    path = generate_docx(
        {
            "title": "Тестовое КП",
            "client_name": "ООО Ромашка",
            "items": [{"name": "Модуль P2.5", "quantity": 24, "unit": "шт", "price": 15000}],
            "total": 360000,
            "notes": "Гарантия 24 месяца",
        }
    )

    document = Document(path)
    full_text = "\n".join(p.text for p in document.paragraphs)

    assert document.paragraphs[0].text == "Тестовое КП"
    assert "ООО Ромашка" in full_text
    assert "Гарантия 24 месяца" in full_text
    assert "Итого: 360000" in full_text

    table = document.tables[0]
    assert table.rows[0].cells[0].text == "Наименование"
    assert table.rows[1].cells[0].text == "Модуль P2.5"


def test_generate_xlsx_creates_file_with_expected_content():
    path = generate_xlsx(
        {
            "title": "Смета по объекту",
            "rows": [
                {"name": "Модуль P2.5", "quantity": 24, "unit_price": 15000},
                {"name": "БП Mean Well", "quantity": 6, "unit_price": 3000},
            ],
        }
    )

    workbook = load_workbook(path)
    sheet = workbook.active

    assert sheet["A1"].value == "Смета по объекту"
    assert sheet.cell(row=3, column=1).value == "Наименование"
    assert sheet.cell(row=4, column=1).value == "Модуль P2.5"
    assert sheet.cell(row=4, column=4).value == 24 * 15000
    assert sheet.cell(row=5, column=4).value == 6 * 3000
    assert sheet.cell(row=6, column=1).value == "Итого"
    assert sheet.cell(row=6, column=4).value == 24 * 15000 + 6 * 3000
