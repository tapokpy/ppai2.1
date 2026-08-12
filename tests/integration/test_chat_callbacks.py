from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from docx import Document
from openpyxl import load_workbook
from sqlalchemy import select

from app.bot.handlers.chat import (
    ask_cloud_callback,
    export_docx_callback,
    export_xlsx_callback,
    save_to_kb_callback,
)
from app.core.database import async_session_maker
from app.models.sqlalchemy.message import Message as MessageModel
from app.models.sqlalchemy.user import User
from tests.integration.conftest import requires_postgres

pytestmark = requires_postgres


async def _seed_user_and_message(telegram_message_id: int = 100) -> tuple[User, MessageModel]:
    async with async_session_maker() as session:
        user = User(telegram_id=777, username="engineer")
        session.add(user)
        await session.commit()
        await session.refresh(user)

        message = MessageModel(
            user_id=user.id,
            telegram_message_id=telegram_message_id,
            prompt="Сколько модулей нужно для экрана 3x2?",
            response="Нужно 24 модуля",
            source="local",
            context_used=False,
        )
        session.add(message)
        await session.commit()

    return user, message


async def _seed_user_and_calculator_message(telegram_message_id: int = 200) -> User:
    async with async_session_maker() as session:
        user = User(telegram_id=778, username="engineer2")
        session.add(user)
        await session.commit()
        await session.refresh(user)

        message = MessageModel(
            user_id=user.id,
            telegram_message_id=telegram_message_id,
            prompt="Расчёт модулей: 3x2 м, шаг 2.5 мм",
            response="Модулей: 24",
            source="calculator",
            context_used=False,
            structured_data={
                "kind": "module_calculation",
                "title": "Коммерческое предложение: модульный экран",
                "items": [{"name": "Модуль дисплея, шаг пикселя 2.5 мм", "quantity": 24, "unit": "шт", "price": ""}],
                "rows": [{"name": "Модуль дисплея, шаг пикселя 2.5 мм", "quantity": 24, "unit_price": 0}],
            },
        )
        session.add(message)
        await session.commit()

    return user


@pytest.mark.asyncio
async def test_export_docx_callback_sends_document(clean_db):
    user, _ = await _seed_user_and_message()
    callback = SimpleNamespace(
        data="export_docx:100",
        message=SimpleNamespace(answer_document=AsyncMock()),
        answer=AsyncMock(),
    )

    await export_docx_callback(callback, user)

    callback.message.answer_document.assert_awaited_once()
    callback.answer.assert_awaited_once()


@pytest.mark.asyncio
async def test_export_xlsx_callback_sends_document(clean_db):
    user, _ = await _seed_user_and_message()
    callback = SimpleNamespace(
        data="export_xlsx:100",
        message=SimpleNamespace(answer_document=AsyncMock()),
        answer=AsyncMock(),
    )

    await export_xlsx_callback(callback, user)

    callback.message.answer_document.assert_awaited_once()


@pytest.mark.asyncio
async def test_export_docx_callback_uses_structured_calculator_data(clean_db):
    user = await _seed_user_and_calculator_message()
    callback = SimpleNamespace(
        data="export_docx:200",
        message=SimpleNamespace(answer_document=AsyncMock()),
        answer=AsyncMock(),
    )

    await export_docx_callback(callback, user)

    sent_file = callback.message.answer_document.call_args.args[0]
    document = Document(sent_file.path)
    table = document.tables[0]

    assert document.paragraphs[0].text == "Коммерческое предложение: модульный экран"
    assert table.rows[1].cells[0].text == "Модуль дисплея, шаг пикселя 2.5 мм"
    assert table.rows[1].cells[1].text == "24"


@pytest.mark.asyncio
async def test_export_xlsx_callback_uses_structured_calculator_data(clean_db):
    user = await _seed_user_and_calculator_message()
    callback = SimpleNamespace(
        data="export_xlsx:200",
        message=SimpleNamespace(answer_document=AsyncMock()),
        answer=AsyncMock(),
    )

    await export_xlsx_callback(callback, user)

    sent_file = callback.message.answer_document.call_args.args[0]
    workbook = load_workbook(sent_file.path)
    sheet = workbook.active

    assert sheet.cell(row=4, column=1).value == "Модуль дисплея, шаг пикселя 2.5 мм"
    assert sheet.cell(row=4, column=2).value == 24


@pytest.mark.asyncio
async def test_export_docx_callback_missing_message_shows_alert(clean_db):
    user, _ = await _seed_user_and_message()
    callback = SimpleNamespace(
        data="export_docx:999",
        message=SimpleNamespace(answer_document=AsyncMock()),
        answer=AsyncMock(),
    )

    await export_docx_callback(callback, user)

    callback.answer.assert_awaited_once_with("Сообщение не найдено", show_alert=True)
    callback.message.answer_document.assert_not_awaited()


@pytest.mark.asyncio
async def test_export_xlsx_callback_missing_message_shows_alert(clean_db):
    user, _ = await _seed_user_and_message()
    callback = SimpleNamespace(
        data="export_xlsx:999",
        message=SimpleNamespace(answer_document=AsyncMock()),
        answer=AsyncMock(),
    )

    await export_xlsx_callback(callback, user)

    callback.answer.assert_awaited_once_with("Сообщение не найдено", show_alert=True)
    callback.message.answer_document.assert_not_awaited()


@pytest.mark.asyncio
async def test_ask_cloud_callback_missing_message_shows_alert(clean_db):
    user, _ = await _seed_user_and_message()
    cascade_router = AsyncMock()
    callback = SimpleNamespace(
        data="ask_cloud:999",
        message=SimpleNamespace(message_id=101, answer=AsyncMock()),
        answer=AsyncMock(),
    )

    await ask_cloud_callback(callback, user, cascade_router)

    callback.answer.assert_awaited_once_with("Сообщение не найдено", show_alert=True)
    cascade_router.process_query.assert_not_awaited()


@pytest.mark.asyncio
async def test_ask_cloud_callback_reprocesses_via_cloud(clean_db):
    user, _ = await _seed_user_and_message()
    cascade_router = AsyncMock()
    cascade_router.process_query.return_value = {
        "text": "Облачный ответ",
        "source": "cloud",
        "context_used": False,
        "elapsed_seconds": 4.1,
    }

    callback = SimpleNamespace(
        data="ask_cloud:100",
        message=SimpleNamespace(message_id=101, answer=AsyncMock()),
        answer=AsyncMock(),
    )

    with patch("app.bot.handlers.admin.settings") as settings_mock:
        settings_mock.admin_ids = []
        await ask_cloud_callback(callback, user, cascade_router)

    cascade_router.process_query.assert_awaited_once_with(
        user_id=user.id,
        prompt="Сколько модулей нужно для экрана 3x2?",
        use_cloud_override=True,
    )
    callback.message.answer.assert_awaited_once()
    # Non-admin: no metrics footer, no export/ask-cloud/save-kb keyboard.
    assert callback.message.answer.call_args.args[0] == "Облачный ответ"
    assert callback.message.answer.call_args.kwargs.get("reply_markup") is None

    async with async_session_maker() as session:
        messages = (await session.execute(select(MessageModel))).scalars().all()

    assert len(messages) == 2
    assert any(m.source == "cloud" for m in messages)


@pytest.mark.asyncio
async def test_save_to_kb_callback_stores_summary_in_rag(clean_db):
    user, _ = await _seed_user_and_message()

    cascade_router = MagicMock()
    cascade_router.local_llm.generate = AsyncMock(return_value="## Краткая инструкция\nОтвет: 24 модуля")
    cascade_router.rag_engine.add_documents = MagicMock()

    callback = SimpleNamespace(data="save_kb:100", answer=AsyncMock())

    await save_to_kb_callback(callback, user, cascade_router)

    cascade_router.local_llm.generate.assert_awaited_once()
    cascade_router.rag_engine.add_documents.assert_called_once()
    call_kwargs = cascade_router.rag_engine.add_documents.call_args.kwargs
    assert call_kwargs["texts"] == ["## Краткая инструкция\nОтвет: 24 модуля"]
    assert call_kwargs["metadatas"][0]["source"] == "harvested"
    callback.answer.assert_awaited_once_with("Инструкция сохранена в базу знаний")


@pytest.mark.asyncio
async def test_save_to_kb_callback_missing_message_shows_alert(clean_db):
    user, _ = await _seed_user_and_message()
    cascade_router = MagicMock()
    callback = SimpleNamespace(data="save_kb:999", answer=AsyncMock())

    await save_to_kb_callback(callback, user, cascade_router)

    callback.answer.assert_awaited_once_with("Сообщение не найдено", show_alert=True)
