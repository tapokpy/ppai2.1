from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

TEMP_DIR = Path("data/temp")


def _ensure_temp_dir() -> Path:
    TEMP_DIR.mkdir(parents=True, exist_ok=True)
    return TEMP_DIR


def generate_docx(proposal_data: dict[str, Any]) -> str:
    document = Document()
    document.add_heading(proposal_data.get("title", "Коммерческое предложение"), level=1)

    client_name = proposal_data.get("client_name")
    if client_name:
        document.add_paragraph(f"Клиент: {client_name}")

    items = proposal_data.get("items", [])
    if items:
        table = document.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for cell, header in zip(table.rows[0].cells, ["Наименование", "Кол-во", "Ед.", "Цена"]):
            cell.text = header

        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("name", ""))
            row_cells[1].text = str(item.get("quantity", ""))
            row_cells[2].text = str(item.get("unit", ""))
            row_cells[3].text = str(item.get("price", ""))

    total = proposal_data.get("total")
    if total is not None:
        run = document.add_paragraph().add_run(f"Итого: {total}")
        run.bold = True
        run.font.size = Pt(12)

    notes = proposal_data.get("notes")
    if notes:
        document.add_paragraph(notes)

    file_path = _ensure_temp_dir() / f"proposal_{uuid4().hex}.docx"
    document.save(file_path)
    return str(file_path)


def generate_xlsx(estimate_data: dict[str, Any]) -> str:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Смета"

    sheet.merge_cells("A1:D1")
    sheet["A1"] = estimate_data.get("title", "Смета")
    sheet["A1"].font = Font(bold=True, size=14)
    sheet["A1"].alignment = Alignment(horizontal="center")

    header_row = 3
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    for col_index, header in enumerate(["Наименование", "Кол-во", "Цена за ед.", "Сумма"], start=1):
        cell = sheet.cell(row=header_row, column=col_index, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    rows = estimate_data.get("rows", [])
    row_index = header_row + 1
    for row in rows:
        quantity = row.get("quantity", 0)
        unit_price = row.get("unit_price", 0)
        sheet.cell(row=row_index, column=1, value=row.get("name", ""))
        sheet.cell(row=row_index, column=2, value=quantity)
        sheet.cell(row=row_index, column=3, value=unit_price)
        sheet.cell(row=row_index, column=4, value=quantity * unit_price)
        row_index += 1

    total_row = row_index
    sheet.cell(row=total_row, column=1, value="Итого").font = Font(bold=True)
    total_value = sum(row.get("quantity", 0) * row.get("unit_price", 0) for row in rows)
    sheet.cell(row=total_row, column=4, value=total_value).font = Font(bold=True)

    for col_index in range(1, 5):
        column_letter = get_column_letter(col_index)
        max_length = max(
            (len(str(sheet.cell(row=r, column=col_index).value or "")) for r in range(1, total_row + 1)),
            default=10,
        )
        sheet.column_dimensions[column_letter].width = max_length + 4

    file_path = _ensure_temp_dir() / f"estimate_{uuid4().hex}.xlsx"
    workbook.save(file_path)
    return str(file_path)
